from __future__ import annotations

import math
import re
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"D:\zuoye\TrafficVisionAnalysis")
OUT = ROOT / "交付文档"
IMG = OUT / "images"
TEX = OUT / "latex"


def ensure_dirs() -> None:
    IMG.mkdir(parents=True, exist_ok=True)
    TEX.mkdir(parents=True, exist_ok=True)


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
    ]
    for p in candidates:
        if Path(p).exists():
            return ImageFont.truetype(p, size=size)
    return ImageFont.load_default()


def wrap_text(text: str, draw: ImageDraw.ImageDraw, fnt: ImageFont.FreeTypeFont, max_width: int) -> list[str]:
    lines: list[str] = []
    for raw in text.split("\n"):
        line = ""
        for ch in raw:
            test = line + ch
            if draw.textbbox((0, 0), test, font=fnt)[2] <= max_width:
                line = test
            else:
                if line:
                    lines.append(line)
                line = ch
        lines.append(line)
    return lines


def draw_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    title: str,
    items: list[str],
    fill: str = "#F7FAFF",
    outline: str = "#3B82F6",
    title_fill: str = "#DBEAFE",
) -> None:
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=18, fill=fill, outline=outline, width=4)
    draw.rounded_rectangle((x1, y1, x2, y1 + 82), radius=18, fill=title_fill, outline=outline, width=0)
    draw.rectangle((x1, y1 + 58, x2, y1 + 82), fill=title_fill)
    title_font = font(42, True)
    body_font = font(35)
    draw.text((x1 + 28, y1 + 17), title, font=title_font, fill="#0B2545")
    y = y1 + 108
    for item in items:
        wrapped = wrap_text(item, draw, body_font, x2 - x1 - 72)
        for i, line in enumerate(wrapped):
            prefix = "• " if i == 0 else "  "
            draw.text((x1 + 28, y), prefix + line, font=body_font, fill="#172033")
            y += 50
        y += 8


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], label: str = "") -> None:
    label = ""
    remap = {
        ((760, 455), (920, 455)): ((930, 560), (1080, 560)),
        ((1620, 500), (1780, 500)): ((1950, 600), (2100, 600)),
        ((2570, 500), (2820, 500)): ((3100, 600), (3260, 600)),
        ((1270, 730), (1270, 970)): ((1515, 900), (1515, 1250)),
        ((2180, 790), (1900, 970)): ((2600, 980), (2200, 1250)),
        ((3160, 680), (2260, 970)): ((3730, 820), (2820, 1250)),
        ((1290, 435), (1350, 435)): ((1320, 475), (1450, 475)),
        ((1910, 630), (1910, 720)): ((2050, 690), (2050, 850)),
        ((1350, 915), (1230, 915)): ((2780, 1065), (2650, 1065)),
        ((670, 1110), (670, 1200)): ((720, 1280), (720, 1280)),
        ((1230, 1395), (1350, 1395)): ((1320, 1065), (1450, 1065)),
        ((1910, 1590), (1910, 1665)): ((2050, 1500), (2050, 1280)),
    }
    start, end = remap.get((start, end), (start, end))
    if start == end:
        return
    sx, sy = start
    ex, ey = end
    draw.line((sx, sy, ex, ey), fill="#334155", width=5)
    ang = math.atan2(ey - sy, ex - sx)
    size = 18
    pts = [
        (ex, ey),
        (ex - size * math.cos(ang - 0.45), ey - size * math.sin(ang - 0.45)),
        (ex - size * math.cos(ang + 0.45), ey - size * math.sin(ang + 0.45)),
    ]
    draw.polygon(pts, fill="#334155")
    if label:
        mid = ((sx + ex) // 2, (sy + ey) // 2)
        f = font(20, True)
        bbox = draw.textbbox((0, 0), label, font=f)
        draw.rounded_rectangle(
            (mid[0] - (bbox[2] - bbox[0]) // 2 - 12, mid[1] - 22, mid[0] + (bbox[2] - bbox[0]) // 2 + 12, mid[1] + 18),
            radius=8,
            fill="#FFFFFF",
            outline="#CBD5E1",
        )
        draw.text((mid[0] - (bbox[2] - bbox[0]) // 2, mid[1] - 16), label, font=f, fill="#1E3A8A")


def generate_technical_architecture() -> Path:
    path = IMG / "technical-architecture-hd.png"
    im = Image.new("RGB", (4300, 2100), "#FFFFFF")
    draw = ImageDraw.Draw(im)
    title_font = font(66, True)
    draw.text((90, 50), "技术架构分层图", font=title_font, fill="#0B2545")
    draw.text((90, 112), "设备采集层 → 边缘服务层 → 算法推理层 → 数据持久层 / 前端展示层", font=font(28), fill="#475569")

    boxes = {
        "device": (100, 300, 930, 820),
        "edge": (1080, 300, 1950, 900),
        "algo": (2100, 300, 3100, 980),
        "data": (3260, 300, 4200, 820),
        "front": (900, 1250, 3400, 1870),
    }
    draw_box(draw, boxes["device"], "设备采集层", ["手机 IP Webcam：HTTP MJPEG / RTSP", "网络摄像头：RTSP / MJPEG", "本地沙盘视频：MP4 文件"], "#F8FBFF", "#2563EB", "#DBEAFE")
    draw_box(draw, boxes["edge"], "边缘服务层（笔记本）", ["FastAPI 后端服务", "VideoStreamService：OpenCV 拉流 / MJPEG 转发", "DashboardState：实时状态聚合", "模型与参数管理：阈值 / 模型切换 / Demo 模式"], "#F7FEFF", "#0891B2", "#CFFAFE")
    draw_box(draw, boxes["algo"], "算法推理层", ["YOLOv11s-visdrone：车辆 / 障碍物基础检测", "ByteTrack：多目标跟踪", "HyperLPR3 / 电子 ID：车牌与身份识别", "S2M：道路未知异物分割增强", "规则引擎：拥堵 / 禁停 / 闸机决策"], "#F8FFF9", "#16A34A", "#DCFCE7")
    draw_box(draw, boxes["data"], "数据持久层", ["SQLite / JSONL", "设备数据、交通统计", "事件告警、通行记录", "模型配置、白名单信息"], "#FFFBEB", "#D97706", "#FEF3C7")
    draw_box(draw, boxes["front"], "前端展示层", ["React + Vite 浅色沙盘大屏", "沙盘全景 MJPEG 画面", "车流 / 拥堵 / 设备状态", "视频源 / 模型 / 闸机控制", "事件告警 / 运行日志"], "#FCF7FF", "#7C3AED", "#EDE9FE")

    arrow(draw, (760, 455), (920, 455), "视频流")
    arrow(draw, (1620, 500), (1780, 500), "最新帧")
    arrow(draw, (2570, 500), (2820, 500), "结果落库")
    arrow(draw, (1270, 730), (1270, 970), "MJPEG / REST")
    arrow(draw, (2180, 790), (1900, 970), "检测与决策")
    arrow(draw, (3160, 680), (2260, 970), "历史数据")
    im.save(path, quality=98)
    return path


def generate_functional_architecture() -> Path:
    path = IMG / "functional-architecture-hd.png"
    im = Image.new("RGB", (4200, 2300), "#FFFFFF")
    draw = ImageDraw.Draw(im)
    draw.text((90, 50), "功能架构分层图", font=font(48, True), fill="#0B2545")
    draw.text((90, 112), "底层视频接入支撑上层视觉感知、交通分析、规则决策、数据管理和可视化展示", font=font(27), fill="#475569")

    layers = [
        ("系统运维层", ["管理边缘设备", "系统状态监控", "视频源切换", "模型状态查看"], "#F8FAFC", "#64748B"),
        ("可视化展示层", ["浅色沙盘大屏", "实时监控画面", "检测框与车辆 ID", "拥堵热力与 KPI", "事件告警与运行日志"], "#FAF5FF", "#7C3AED"),
        ("数据管理层", ["保存统计数据到数据库", "查询历史统计与事件记录", "导出分析报告", "维护白名单车辆", "管理模型与检测参数"], "#FFF7ED", "#EA580C"),
        ("规则决策层", ["禁停区域告警", "道路异常告警", "闸机白名单决策", "事件等级判定", "验收演示流程"], "#FEF2F2", "#DC2626"),
        ("交通分析层", ["车流量统计", "车速估计", "拥堵等级判定", "车流密度热力图", "历史趋势统计"], "#F0FDF4", "#16A34A"),
        ("视觉感知层", ["YOLOv11 车辆检测", "ByteTrack 车辆跟踪", "车牌 / 电子 ID 识别", "障碍物识别", "S2M 未知异物分割增强"], "#EFF6FF", "#2563EB"),
        ("视频接入层", ["接入手机视频流", "接入 RTSP/MJPEG 网络摄像头", "接入本地沙盘视频段", "视频状态监测与重连", "图像预处理：缩放 / ROI / 增强"], "#ECFEFF", "#0891B2"),
    ]
    left_x, mid_x, right_x, w, h = 120, 1450, 2780, 1200, 430
    positions = [
        (left_x, 260), (mid_x, 260),
        (right_x, 260), (left_x, 850),
        (mid_x, 850), (right_x, 850),
        (1450, 1500),
    ]
    for idx, ((title, items, fill, outline), (x, y)) in enumerate(zip(layers, positions)):
        draw_box(draw, (x, y, x + w, y + h), title, items, fill, outline, "#FFFFFF")
    # Clear, low-clutter layer flow.
    arrow(draw, (1290, 435), (1350, 435), "协同")
    arrow(draw, (1910, 630), (1910, 720), "支撑")
    arrow(draw, (1350, 915), (1230, 915), "数据")
    arrow(draw, (670, 1110), (670, 1200), "规则")
    arrow(draw, (1230, 1395), (1350, 1395), "分析")
    arrow(draw, (1910, 1590), (1910, 1665), "感知")
    im.save(path, quality=98)
    return path


def generate_er_diagram() -> Path:
    path = IMG / "database-er-hd.png"
    im = Image.new("RGB", (3000, 1900), "#FFFFFF")
    draw = ImageDraw.Draw(im)
    draw.text((90, 45), "数据库 ER 图", font=font(48, True), fill="#0B2545")
    draw.text((90, 108), "核心实体：设备、模型配置、交通统计、事件记录、通行记录、白名单车辆", font=font(28), fill="#475569")

    def table(x, y, title, fields, color):
        width, row_h = 660, 54
        height = 70 + row_h * len(fields)
        draw.rounded_rectangle((x, y, x + width, y + height), radius=14, fill="#FFFFFF", outline=color, width=4)
        draw.rounded_rectangle((x, y, x + width, y + 70), radius=14, fill=color, outline=color)
        draw.rectangle((x, y + 52, x + width, y + 70), fill=color)
        draw.text((x + 22, y + 18), title, font=font(26, True), fill="#FFFFFF")
        f = font(21)
        for i, field in enumerate(fields):
            yy = y + 70 + i * row_h
            fill = "#F8FAFC" if i % 2 == 0 else "#FFFFFF"
            draw.rectangle((x, yy, x + width, yy + row_h), fill=fill, outline="#CBD5E1")
            draw.text((x + 20, yy + 13), field, font=f, fill="#172033")
        return (x, y, x + width, y + height)

    t_devices = table(1200, 250, "DEVICES", ["device_id PK", "name", "type", "stream_url", "location", "status", "last_seen"], "#2563EB")
    t_models = table(120, 250, "MODEL_CONFIGS", ["model_id PK", "model_name", "model_path", "confidence", "iou", "inference_size", "enabled"], "#7C3AED")
    t_stats = table(120, 1050, "TRAFFIC_STATS", ["id PK", "device_id FK", "model_id FK", "timestamp", "vehicle_count", "density", "avg_speed", "congestion_level"], "#16A34A")
    t_events = table(1200, 1050, "EVENTS", ["event_id PK", "device_id FK", "model_id FK", "type", "severity", "description", "created_at"], "#DC2626")
    t_white = table(2220, 250, "WHITELIST", ["identity PK", "owner", "vehicle_type", "status", "updated_at"], "#D97706")
    t_plate = table(2220, 1050, "PLATE_RECORDS", ["id PK", "device_id FK", "whitelist_identity FK", "plate_no", "electronic_id", "gate_action", "confidence", "timestamp"], "#0891B2")

    arrow(draw, (1200, 570), (780, 1170), "1:N")
    arrow(draw, (1860, 570), (1530, 1050), "1:N")
    arrow(draw, (1860, 600), (2220, 1180), "1:N")
    arrow(draw, (780, 570), (420, 1050), "1:N")
    arrow(draw, (780, 600), (1200, 1200), "1:N")
    arrow(draw, (2550, 640), (2550, 1050), "1:N")
    im.save(path, quality=98)
    return path


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(9.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill.replace("#", ""))
    tc_pr.append(shd)


def set_cell_borders(cell, color: str = "D0D7DE") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def style_table(tbl) -> None:
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_idx, row in enumerate(tbl.rows):
        for cell in row.cells:
            set_cell_borders(cell)
            if row_idx == 0:
                shade_cell(cell, "E8EEF5")
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True


def setup_doc(title: str, subtitle: str) -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.82)
    sec.bottom_margin = Inches(0.78)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [("Heading 1", 16, "2E74B5"), ("Heading 2", 13, "2E74B5"), ("Heading 3", 12, "1F4D78")]:
        st = styles[name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = True
        st.paragraph_format.space_before = Pt(12 if name != "Heading 1" else 16)
        st.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(24)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.color.rgb = RGBColor(11, 37, 69)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(subtitle)
    r2.font.size = Pt(12)
    r2.font.color.rgb = RGBColor(71, 85, 105)
    doc.add_paragraph()
    return doc


def add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False
    tbl.columns[0].width = Cm(4)
    tbl.columns[1].width = Cm(12)
    set_cell_text(tbl.rows[0].cells[0], "项目", True)
    set_cell_text(tbl.rows[0].cells[1], "内容", True)
    for k, v in rows:
        cells = tbl.add_row().cells
        set_cell_text(cells[0], k, True)
        set_cell_text(cells[1], v)
    style_table(tbl)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.autofit = False
    for i, h in enumerate(headers):
        set_cell_text(tbl.rows[0].cells[i], h, True)
        if widths:
            tbl.columns[i].width = Cm(widths[i])
    for row in rows:
        cells = tbl.add_row().cells
        for i, v in enumerate(row):
            set_cell_text(cells[i], v)
    style_table(tbl)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.65)
        p.add_run(item)


def add_picture(doc: Document, path: Path, caption: str, width_cm: float = 16.0) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cap.runs:
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(71, 85, 105)


def build_requirement_doc(use_case_img: Path) -> Path:
    doc = setup_doc("云边端协同智慧交通视觉感知系统", "需求分析报告 V1.0")
    add_kv_table(doc, [
        ("项目场景", "校园 704 智慧交通沙盘"),
        ("系统形态", "手机/网络摄像头/本地视频采集 + 笔记本分析服务 + 算法推理 + Web 大屏展示"),
        ("当前版本重点", "视频接入、浅色沙盘大屏、YOLOv11 检测、ByteTrack 跟踪预留、S2M 障碍物增强方案"),
    ])
    doc.add_heading("1. 项目目标", level=1)
    doc.add_paragraph("本系统面向智慧交通沙盘演示场景，构建从视频采集、边缘接入、车辆检测、交通统计、异常告警到前端展示的闭环。系统优先保证验收演示稳定，同时为后续算法训练、多摄像头融合和数据持久化扩展保留接口。")
    doc.add_heading("2. 用户角色", level=1)
    add_table(doc, ["角色", "职责"], [
        ["交通监控员", "接入视频流，查看实时监控画面、车辆检测、拥堵热力、异常告警和闸机通行决策。"],
        ["系统管理员", "维护边缘设备、模型参数、白名单车辆、历史统计和事件数据。"],
        ["算法开发人员", "接入 YOLOv11、ByteTrack、车牌识别和 S2M 障碍物增强等算法能力。"],
    ], [4, 12])
    doc.add_heading("3. 功能需求概览", level=1)
    add_table(doc, ["编号", "功能", "优先级", "说明"], [
        ["F01", "视频流接入", "P0", "支持手机 IP Webcam、RTSP/MJPEG 网络摄像头、本地视频段和电脑摄像头。"],
        ["F02", "车辆检测与跟踪", "P0", "以 YOLOv11 识别车辆，ByteTrack 维护车辆 ID。"],
        ["F09", "障碍物识别", "P1", "基础模式复用 YOLOv11 检测结果，S2M 作为未知异物分割增强模块。"],
        ["F10", "车流密度热力图", "P1", "基于车辆位置生成拥堵热力与热点区域。"],
        ["F11", "Web 实时监控页面", "P0", "采用浅色沙盘大屏样式，中心展示 MJPEG 实时画面，左右展示状态、控制和告警。"],
        ["F14", "历史统计查询", "P1", "查询车流、拥堵、事件和通行记录。"],
        ["F17", "设备与模型管理", "P1", "管理视频源、模型状态、检测参数和白名单数据。"],
    ], [1.5, 4, 2, 8.5])
    doc.add_heading("4. 用例图与核心用例", level=1)
    add_picture(doc, use_case_img, "图 1 用例图：云边端协同智慧交通视觉感知系统", 16.5)
    add_table(doc, ["用例", "参与者", "主要结果"], [
        ["接入手机视频流", "交通监控员", "系统接入 MJPEG/RTSP 视频并显示实时画面。"],
        ["查看车辆检测与跟踪结果", "交通监控员", "画面叠加车辆框、类别、置信度和跟踪 ID。"],
        ["查看车流统计与拥堵等级", "交通监控员", "展示当前车辆数、密度、平均速度和拥堵等级。"],
        ["查看禁停/道路异常告警", "交通监控员", "展示禁停、道路异物、障碍物或拥堵告警。"],
        ["管理模型与检测参数", "系统管理员/算法开发人员", "切换模型并调整置信度、检测间隔、ROI 等参数。"],
        ["接入算法推理服务", "算法开发人员", "后端调用算法接口并接收检测、跟踪、识别结果。"],
    ], [4.3, 4.3, 7.4])
    doc.add_heading("5. 非功能需求", level=1)
    add_table(doc, ["类别", "要求"], [
        ["性能", "端到端画面延迟小于 1 秒；基础检测链路目标处理帧率不低于 10 FPS。"],
        ["可靠性", "视频断连后显示离线状态并支持重新接入；算法不可用时降级为视频监控和演示数据。"],
        ["可用性", "页面采用浅色高对比度大屏布局，支持 1920×1080 展示和窄窗口纵向堆叠。"],
        ["可维护性", "前端、后端、算法模块解耦，通过 REST/MJPEG/HTTP 推理接口协作。"],
    ], [4, 12])
    doc.add_heading("6. 数据与接口约束", level=1)
    add_bullets(doc, [
        "视频输入优先使用手机热点、USB 网络共享或自建局域网，避免校园网客户端隔离导致无法传输。",
        "前端当前采用 MJPEG 视频流和 REST 轮询分析结果，不再以 WebSocket 推帧作为主方案。",
        "S2M 仅作为障碍物识别增强模块，不作为第一版实时主链路强依赖。",
        "历史统计、事件记录、通行记录和模型配置可优先使用 SQLite 或 JSONL 落盘。"
    ])
    path = OUT / "需求分析报告.docx"
    doc.save(path)
    return path


def build_design_doc(tech_img: Path, func_img: Path, er_img: Path) -> Path:
    doc = setup_doc("云边端协同智慧交通视觉感知系统", "系统设计报告 V1.0")
    add_kv_table(doc, [
        ("总体架构", "端侧采集、边缘服务、算法推理、数据持久、前端展示五层架构"),
        ("前端实现", "React + Vite 浅色沙盘大屏，MJPEG 主画面 + REST 轮询状态"),
        ("算法实现", "YOLOv11s-visdrone、ByteTrack、HyperLPR3、S2M 增强模块"),
    ])
    doc.add_heading("1. 设计目标与原则", level=1)
    doc.add_paragraph("系统采用分层视频分析设计，以笔记本作为分析服务器，支持手机、RTSP/MJPEG 网络摄像头、本地视频段等多种视频源接入。设计重点是保证演示闭环稳定、接口清晰、算法可替换、前端展示清楚。")
    add_bullets(doc, [
        "模块解耦：前端、后端、算法三部分独立开发，通过 REST API、MJPEG 视频流和算法 HTTP 接口通信。",
        "演示优先：保留本地视频和 Demo 模式，降低现场网络不稳定对验收的影响。",
        "算法可扩展：YOLOv11 作为车辆检测主模型，ByteTrack、HyperLPR3 和 S2M 按模块接入。",
    ])
    doc.add_heading("2. 技术架构分层图", level=1)
    add_picture(doc, tech_img, "图 1 技术架构分层图", 17.0)
    doc.add_heading("3. 功能架构分层图", level=1)
    add_picture(doc, func_img, "图 2 功能架构分层图", 12.0)
    doc.add_heading("4. 数据库 ER 图", level=1)
    add_picture(doc, er_img, "图 3 数据库 ER 图", 17.0)
    doc.add_heading("5. 核心接口设计", level=1)
    add_table(doc, ["接口", "方法", "路径", "说明"], [
        ["启动视频流", "POST", "/api/video/start", "传入视频源地址，开始拉流。"],
        ["视频 MJPEG", "GET", "/api/video/mjpeg", "向前端输出 multipart/x-mixed-replace 视频流。"],
        ["视频状态", "GET", "/api/video/status", "返回连接状态、帧率、分辨率和当前源。"],
        ["最新分析", "GET", "/api/analysis/latest", "返回检测框、车辆统计、拥堵等级和事件摘要。"],
        ["模型列表", "GET", "/api/models", "查询可用模型和当前模型。"],
        ["闸机决策", "POST", "/api/whitelist/decision", "根据车牌或电子 ID 返回放行/拒绝结果。"],
    ], [3.2, 2, 4.2, 6.6])
    doc.add_heading("6. 前端设计", level=1)
    doc.add_paragraph("前端采用浅色沙盘大屏样式，主监控画面位于中心，左侧展示车流统计、拥堵热力和设备状态，右侧集中视频源接入、模型状态、闸机决策和事件告警，底部显示运行日志。")
    doc.add_heading("7. 算法设计", level=1)
    add_table(doc, ["模块", "方案", "说明"], [
        ["车辆检测", "YOLOv11s-visdrone", "适配沙盘小目标和俯视视角车辆检测。"],
        ["多目标跟踪", "ByteTrack", "维护车辆轨迹和 track_id，用于计数和速度估计。"],
        ["车牌识别", "HyperLPR3 / 电子 ID", "输出车牌号或电子 ID，服务于闸机白名单决策。"],
        ["障碍物识别", "YOLOv11 基础过滤 + S2M 增强", "基础模式实时，S2M 用于未知异物和不规则障碍分割。"],
        ["交通分析", "ROI 计数 + 热力图 + 规则引擎", "生成车流统计、拥堵等级、禁停和道路异常告警。"],
    ], [4, 4.2, 7.8])
    doc.add_heading("8. 部署与运行", level=1)
    add_bullets(doc, [
        "开发环境：后端 localhost:8000，前端 localhost:5173，算法服务可本机或 GPU 云部署。",
        "视频源：推荐手机热点、USB 网络共享或自建路由器局域网；校园网可能存在客户端隔离。",
        "数据落盘：第一版优先 SQLite/JSONL，后续可迁移到 MySQL 或 PostgreSQL。",
    ])
    path = OUT / "系统设计报告.docx"
    doc.save(path)
    return path


def latex_escape(s: str) -> str:
    return (
        s.replace("\\", r"\textbackslash{}")
        .replace("&", r"\&")
        .replace("%", r"\%")
        .replace("$", r"\$")
        .replace("#", r"\#")
        .replace("_", r"\_")
        .replace("{", r"\{")
        .replace("}", r"\}")
    )


def write_latex_sources() -> None:
    base = r"""\documentclass[UTF8,12pt]{ctexart}
\usepackage{geometry}
\usepackage{graphicx}
\usepackage{booktabs}
\usepackage{longtable}
\usepackage{xcolor}
\geometry{a4paper,margin=2.2cm}
\definecolor{mainblue}{HTML}{2E74B5}
\title{%s}
\author{云边端协同智慧交通视觉感知系统}
\date{2026年7月}
\begin{document}
\maketitle
%s
\end{document}
"""
    req_body = r"""
\section{项目目标}
本报告面向校园智慧交通沙盘项目，定义视频接入、车辆检测、交通统计、异常告警、闸机决策和前端展示等核心需求。
\section{用例图}
\includegraphics[width=\linewidth]{../docs/images/use-case-diagram.png}
\section{核心需求}
系统采用 MJPEG 视频流和 REST 轮询分析结果，前端为浅色沙盘大屏。车辆检测主方案为 YOLOv11，障碍物识别采用 YOLOv11 基础过滤，并以 S2M 作为增强分割模块。
"""
    design_body = r"""
\section{技术架构分层图}
\includegraphics[width=\linewidth]{images/technical-architecture-hd.png}
\section{功能架构分层图}
\includegraphics[width=0.75\linewidth]{images/functional-architecture-hd.png}
\section{数据库 ER 图}
\includegraphics[width=\linewidth]{images/database-er-hd.png}
"""
    (TEX / "需求分析报告.tex").write_text(base % ("需求分析报告", req_body), encoding="utf-8")
    (TEX / "系统设计报告.tex").write_text(base % ("系统设计报告", design_body), encoding="utf-8")


def main() -> None:
    ensure_dirs()
    tech = generate_technical_architecture()
    func = generate_functional_architecture()
    er = generate_er_diagram()
    use_case = ROOT / "docs" / "images" / "use-case-diagram.png"
    req_doc = build_requirement_doc(use_case)
    design_doc = build_design_doc(tech, func, er)
    write_latex_sources()
    print(req_doc)
    print(design_doc)


if __name__ == "__main__":
    main()
