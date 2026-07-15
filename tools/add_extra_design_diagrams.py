from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"D:\zuoye\TrafficVisionAnalysis")
OUT = ROOT / "交付文档"
IMG = OUT / "images"
DOCX = OUT / "系统设计报告.docx"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
    ]
    for item in candidates:
        path = Path(item)
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def box(draw, xy, title, body=None, fill="#F8FAFC", stroke="#2563EB"):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=22, fill=fill, outline=stroke, width=5)
    draw.rounded_rectangle((x1, y1, x2, y1 + 82), radius=22, fill="#EAF2FF", outline=stroke, width=0)
    draw.rectangle((x1, y1 + 58, x2, y1 + 82), fill="#EAF2FF")
    draw.text((x1 + 26, y1 + 18), title, font=font(40, True), fill="#0B2545")
    if body:
        y = y1 + 112
        for line in body:
            draw.text((x1 + 30, y), "• " + line, font=font(31), fill="#172033")
            y += 48


def arrow(draw, start, end):
    sx, sy = start
    ex, ey = end
    draw.line((sx, sy, ex, ey), fill="#334155", width=6)
    import math

    angle = math.atan2(ey - sy, ex - sx)
    size = 22
    points = [
        (ex, ey),
        (ex - size * math.cos(angle - 0.45), ey - size * math.sin(angle - 0.45)),
        (ex - size * math.cos(angle + 0.45), ey - size * math.sin(angle + 0.45)),
    ]
    draw.polygon(points, fill="#334155")


def title(draw, text, subtitle=""):
    draw.text((90, 52), text, font=font(66, True), fill="#0B2545")
    if subtitle:
        draw.text((90, 132), subtitle, font=font(32), fill="#475569")


def save(im, name):
    path = IMG / name
    im.save(path, quality=98)
    return path


def system_flow():
    im = Image.new("RGB", (4200, 1600), "white")
    draw = ImageDraw.Draw(im)
    title(draw, "系统运行流程图", "从视频输入到前端展示的端到端处理链路")
    y = 430
    w, h, gap = 520, 330, 75
    nodes = [
        ("视频源输入", ["手机 IP Webcam", "RTSP/MJPEG 摄像头", "本地沙盘视频"]),
        ("后端拉流", ["OpenCV 解码", "MJPEG 转发", "状态监测"]),
        ("图像预处理", ["缩放", "ROI 裁剪", "亮度增强"]),
        ("算法推理", ["YOLOv11 检测", "ByteTrack 跟踪", "车牌 / ID 识别"]),
        ("统计告警", ["车流密度", "拥堵等级", "异常事件"]),
        ("前端大屏", ["实时画面", "指标面板", "事件日志"]),
    ]
    colors = ["#EFF6FF", "#ECFEFF", "#F8FAFC", "#F0FDF4", "#FFF7ED", "#FAF5FF"]
    strokes = ["#2563EB", "#0891B2", "#64748B", "#16A34A", "#EA580C", "#7C3AED"]
    x = 120
    for i, (name, lines) in enumerate(nodes):
        box(draw, (x, y, x + w, y + h), name, lines, colors[i], strokes[i])
        if i < len(nodes) - 1:
            arrow(draw, (x + w, y + h // 2), (x + w + gap, y + h // 2))
        x += w + gap
    return save(im, "system-runtime-flow.png")


def api_flow():
    im = Image.new("RGB", (4200, 1900), "white")
    draw = ImageDraw.Draw(im)
    title(draw, "前后端接口调用流程图", "前端启动视频源后，通过 MJPEG 获取画面，通过 REST 获取状态和分析结果")
    actors = [
        ("前端 React 大屏", 240),
        ("FastAPI 后端", 1420),
        ("VideoStreamService", 2600),
        ("算法推理模块", 3460),
    ]
    for name, x in actors:
        box(draw, (x, 280, x + 560, 430), name, [], "#F8FAFC", "#334155")
        draw.line((x + 280, 430, x + 280, 1660), fill="#CBD5E1", width=4)
    calls = [
        (520, 680, 1700, "POST /api/video/start"),
        (1700, 860, 2880, "open source / read frames"),
        (520, 1040, 1700, "GET /api/video/mjpeg"),
        (1700, 1220, 3740, "frame -> detect / track"),
        (3740, 1380, 1700, "analysis result"),
        (1700, 1540, 520, "GET /api/analysis/latest"),
    ]
    for x1, y, x2, label in calls:
        arrow(draw, (x1, y), (x2, y))
        bx1 = min(x1, x2) + abs(x2 - x1) // 2 - 250
        draw.rounded_rectangle((bx1, y - 42, bx1 + 500, y + 22), radius=12, fill="#FFFFFF", outline="#CBD5E1", width=3)
        draw.text((bx1 + 20, y - 30), label, font=font(28, True), fill="#1E3A8A")
    return save(im, "frontend-backend-api-flow.png")


def frontend_layout():
    im = Image.new("RGB", (4200, 2200), "white")
    draw = ImageDraw.Draw(im)
    title(draw, "前端页面布局示意图", "浅色沙盘大屏：中心视频画面，两侧展示状态、控制、统计和告警信息")
    draw.rounded_rectangle((180, 250, 4020, 2050), radius=34, fill="#F8FBFF", outline="#CBD5E1", width=5)
    box(draw, (320, 380, 1000, 800), "视频源控制", ["手机 MJPEG / RTSP", "本地视频", "模型启动 / 停止"], "#EFF6FF", "#2563EB")
    box(draw, (320, 900, 1000, 1320), "运行状态", ["连接状态", "FPS / 分辨率", "处理延迟"], "#ECFEFF", "#0891B2")
    box(draw, (320, 1420, 1000, 1870), "统计指标", ["车辆数", "拥堵等级", "平均速度"], "#F0FDF4", "#16A34A")
    draw.rounded_rectangle((1120, 380, 3060, 1720), radius=24, fill="#EEF2F7", outline="#334155", width=5)
    draw.text((1860, 980), "沙盘实时视频画面", font=font(56, True), fill="#0B2545")
    draw.text((1780, 1060), "车辆框 / ID / 告警叠加", font=font(38), fill="#475569")
    box(draw, (3180, 380, 3860, 820), "检测结果", ["车辆检测", "ByteTrack ID", "车牌 / 电子 ID"], "#FAF5FF", "#7C3AED")
    box(draw, (3180, 920, 3860, 1360), "事件告警", ["禁停", "道路异常", "障碍物 / 拥堵"], "#FEF2F2", "#DC2626")
    box(draw, (3180, 1460, 3860, 1870), "运行日志", ["视频重连", "模型切换", "接口调用"], "#FFF7ED", "#EA580C")
    return save(im, "frontend-screen-layout.png")


def deployment():
    im = Image.new("RGB", (4200, 1800), "white")
    draw = ImageDraw.Draw(im)
    title(draw, "部署结构图", "手机/网络摄像头、笔记本分析服务、前端浏览器和本地数据文件之间的连接关系")
    box(draw, (140, 430, 760, 820), "手机摄像头", ["IP Webcam", "HTTP MJPEG / RTSP"], "#EFF6FF", "#2563EB")
    box(draw, (140, 980, 760, 1370), "网络摄像头", ["RTSP / MJPEG", "沙盘固定视角"], "#ECFEFF", "#0891B2")
    box(draw, (1180, 500, 2050, 1320), "笔记本边缘服务", ["FastAPI 后端", "OpenCV 拉流", "YOLOv11 / ByteTrack", "S2M 增强预留"], "#F0FDF4", "#16A34A")
    box(draw, (2500, 420, 3200, 820), "前端浏览器", ["React + Vite", "浅色沙盘大屏", "MJPEG + REST"], "#FAF5FF", "#7C3AED")
    box(draw, (2500, 980, 3200, 1380), "本地数据", ["SQLite / JSONL", "统计 / 事件 / 配置"], "#FFF7ED", "#EA580C")
    box(draw, (3500, 700, 4060, 1120), "验收展示屏", ["实时监控", "数据看板", "告警展示"], "#F8FAFC", "#64748B")
    arrow(draw, (760, 625), (1180, 760))
    arrow(draw, (760, 1175), (1180, 1060))
    arrow(draw, (2050, 740), (2500, 620))
    arrow(draw, (2050, 1080), (2500, 1180))
    arrow(draw, (3200, 620), (3500, 850))
    return save(im, "deployment-structure.png")


def add_picture(doc, path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(24.8))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(71, 85, 105)


def append_to_docx(paths):
    doc = Document(DOCX)
    doc.add_page_break()
    doc.add_heading("补充设计图", level=1)
    captions = [
        "图 4 系统运行流程图",
        "图 5 前后端接口调用流程图",
        "图 6 前端页面布局示意图",
        "图 7 部署结构图",
    ]
    for path, caption in zip(paths, captions):
        doc.add_heading(caption.split(" ", 2)[-1], level=2)
        add_picture(doc, path, caption)
    doc.save(DOCX)


def main():
    IMG.mkdir(parents=True, exist_ok=True)
    paths = [system_flow(), api_flow(), frontend_layout(), deployment()]
    append_to_docx(paths)
    print(DOCX)


if __name__ == "__main__":
    main()
